import docx
#参考链接：https://blog.csdn.net/u014465934/article/details/79574206
 
doc = docx.Document()
#用add_heading()将添加一个段落，并使用一种标题样式。add_heading()的参数，是一个标题文本的字符串，以及一个从0到4的整数。
#整数0表示标题时Title样式，这用于文档的顶部。整数1到4是不同标题层次，1是主要的标题，4是最底层的子标题。

doc.add_heading('Header 0',0)
doc.add_heading('Header 1',1)
doc.add_heading('Header 2',2)
doc.add_heading('Header 3',3)
doc.add_heading('Header 4',4)

doc.add_paragraph('Hello world!') 
 
paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
 
#用add_run()方法，向已有段落的末尾添加文本。 
paraObj1.add_run('This text is being added to the second paragraph.') 
 
#换行符，在第一行后面添加换行符
doc.paragraphs[0].runs[0].add_break()
 
doc.add_paragraph('This is on the second page!')
#换页符，在第二行后面添加换页符
doc.paragraphs[1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE) 

#添加图像  
doc.add_picture(r'images\miaomiao.png',width=docx.shared.Inches(2.5),height=docx.shared.Cm(7))

doc.save(r'backup\写入word.docx')
