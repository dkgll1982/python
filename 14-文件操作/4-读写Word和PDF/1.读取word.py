from docx import Document 
import re
#参考链接：https://blog.csdn.net/xtfge0915/article/details/83479922

doc = Document(r'backup\电子表格.docx')
for p in doc.paragraphs:
     if p.style.name=='Normal':                 #读取正文
        print(p.text)

print('='*40)
for p in doc.paragraphs:
    if re.match("^Heading \d+$",p.style.name):  #读取标题
        print(p.text)

print('='*40)

#读取表格
tb = doc.tables
rows=tb[0].rows             #行
cols=rows[0].cells          #列
cell=cols[0]                #单元格
text=cell.text
print('表格第1行第一列：',text)

print('='*40)

index,key,value=[],[],[]
table_index=0
for tb in doc.tables:
 table_index+=1
 row_index=0
 for row in tb.rows:
     row_index+=1
     for cell in row.cells:
         text=""
         for p in cell.paragraphs:##如果cell中有多段，即有回车符
             text+=p.text
         if row_index%2==0:
             value.append(text)#偶数行为值
             index.append(table_index)#这行也可以放在else中
         else:
             key.append(text)#奇数行为属性
print(key)
print(value)
print(index)

print('='*40)
